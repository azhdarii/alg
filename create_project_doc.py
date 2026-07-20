from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBidi = OxmlElement("w:bidi")
    pPr.append(pBidi)

def set_font_for_persian(run, font_name="B Nazanin", size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

def add_heading_rtl(doc, text, level=1):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    if level == 1:
        set_font_for_persian(run, size=16, bold=True, color=(31, 58, 95))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_font_for_persian(run, size=14, bold=True, color=(31, 58, 95))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        set_font_for_persian(run, size=12, bold=True, color=(31, 58, 95))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body_rtl(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_font_for_persian(run, size=12, bold=bold)
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet_rtl(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    bullet_run = p.add_run("• ")
    set_font_for_persian(bullet_run, size=12)
    text_run = p.add_run(text)
    set_font_for_persian(text_run, size=12)
    return p

def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ============================================================
# CREATE DOCUMENT
# ============================================================
doc = Document()
setup_page(doc)

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "B Nazanin"
font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.insert(0, rFonts)
rFonts.set(qn("w:eastAsia"), "B Nazanin")

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("گزارش طراحی اولیه مسئله")
set_font_for_persian(run, size=26, bold=True, color=(31, 58, 95))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("تخصیص دانشجویان به اساتید راهنما")
set_font_for_persian(run, size=22, bold=True, color=(31, 58, 95))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("با استفاده از الگوریتم‌های فراابتکاری")
set_font_for_persian(run, size=18, color=(80, 80, 80))

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("Problem Design Document")
set_font_for_persian(run, size=14, color=(120, 120, 120))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("Multi-Objective Optimization for Student-Supervisor Assignment")
set_font_for_persian(run, size=12, color=(120, 120, 120))

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS PLACEHOLDER
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("فهرست مطالب")
set_font_for_persian(run, size=18, bold=True, color=(31, 58, 95))

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "separate")
fldChar3 = OxmlElement("w:t")
fldChar3.text = "برای بروزرسانی فهرست، کلیک راست کرده و Update Field را انتخاب کنید."
fldChar4 = OxmlElement("w:fldChar")
fldChar4.set(qn("w:fldCharType"), "end")
for x in (fldChar1, instrText, fldChar2, fldChar3, fldChar4):
    run._r.append(x)

doc.add_page_break()

# ============================================================
# SECTION 1: PROBLEM DEFINITION
# ============================================================
add_heading_rtl(doc, "۱. تعریف مسئله", level=1)

add_body_rtl(doc, 
    "هدف این پروژه طراحی یک سیستم هوشمند برای تخصیص دانشجویان کارشناسی ارشد به اساتید راهنما است. "
    "در این مسئله، برای هر دانشجو باید دقیقاً یک استاد راهنما انتخاب شود، به گونه‌ای که علاوه بر تأمین "
    "رضایت دانشجویان و اساتید، توزیع دانشجویان با کیفیت بین اساتید نیز عادلانه باشد.")

add_body_rtl(doc,
    "این مسئله در دسته مسائل بهینه‌سازی ترکیبی (Combinatorial Optimization Problems) قرار می‌گیرد، "
    "زیرا تعداد حالات ممکن تخصیص با افزایش تعداد دانشجویان و اساتید به صورت نمایی رشد می‌کند. "
    "به عنوان مثال، اگر ۶۰ دانشجو و ۱۲ استاد وجود داشته باشد، تعداد حالت‌های ممکن تقریباً برابر است با:")

add_formula(doc, "12^60")

add_body_rtl(doc,
    "که عددی بسیار بزرگ بوده و بررسی تمام حالت‌ها (Exhaustive Search) از نظر زمانی غیرممکن است. "
    "بنابراین استفاده از روش‌های فراابتکاری مانند NSGA-II و Weighted Sum Genetic Algorithm "
    "انتخاب مناسبی برای حل این مسئله خواهد بود.")

add_separator(doc)

# ============================================================
# SECTION 2: ASSUMPTIONS AND DATA
# ============================================================
add_heading_rtl(doc, "۲. پیش‌فرض‌ها و داده‌های مسئله", level=1)

add_body_rtl(doc, "برای مدل‌سازی مسئله، فرضیات زیر در نظر گرفته می‌شوند:")

assumptions = [
    "تعداد دانشجویان و اساتید از قبل مشخص است.",
    "هر دانشجو دقیقاً به یک استاد تخصیص داده می‌شود.",
    "هر استاد می‌تواند چندین دانشجو داشته باشد.",
    "ظرفیت هر استاد از قبل مشخص است.",
    "اولویت‌های دانشجویان نسبت به اساتید قبل از اجرای الگوریتم جمع‌آوری شده است.",
    "اولویت یا امتیاز اساتید نسبت به دانشجویان نیز از قبل مشخص است.",
    "کیفیت هر دانشجو قبل از اجرای الگوریتم محاسبه شده و در طول اجرای الگوریتم ثابت باقی می‌ماند.",
    "هر دانشجو دارای یک گرایش اصلی است.",
    "هر استاد دارای یک گرایش اصلی و مجموعه‌ای از علایق پژوهشی است.",
    "در نسخه اولیه مدل، محدودیت ظرفیت و تطابق گرایش به صورت محدودیت نرم (Soft Constraint) در نظر گرفته می‌شوند.",
]

for item in assumptions:
    add_bullet_rtl(doc, item)

add_heading_rtl(doc, "داده‌های مورد نیاز", level=2)

add_heading_rtl(doc, "اطلاعات دانشجویان", level=3)
student_data = [
    "شناسه دانشجو",
    "نام",
    "گرایش",
    "معدل",
    "دانشگاه محل تحصیل قبلی",
    "امتیاز کیفیت (Quality Score)",
    "علایق پژوهشی",
    "لیست اولویت اساتید",
]
for item in student_data:
    add_bullet_rtl(doc, item)

add_heading_rtl(doc, "اطلاعات اساتید", level=3)
professor_data = [
    "شناسه",
    "نام",
    "گرایش",
    "ظرفیت",
    "مرتبه علمی",
    "علایق پژوهشی",
    "لیست اولویت دانشجویان",
]
for item in professor_data:
    add_bullet_rtl(doc, item)

add_separator(doc)

# ============================================================
# SECTION 3: SOLUTION REPRESENTATION
# ============================================================
add_heading_rtl(doc, "۳. نمایش جواب (Solution Representation)", level=1)

add_body_rtl(doc,
    "یکی از مهم‌ترین مراحل طراحی الگوریتم‌های فراابتکاری، انتخاب نحوه نمایش یک جواب است. "
    "در این پروژه دو نمایش اصلی قابل استفاده هستند.")

add_heading_rtl(doc, "نمایش ماتریسی (Binary Matrix)", level=2)
add_body_rtl(doc,
    "در این روش، یک ماتریس به ابعاد تعداد دانشجویان × تعداد اساتید ساخته می‌شود که مقدار هر خانه صفر یا یک است.")

add_body_rtl(doc, "مزایا:", bold=True)
add_bullet_rtl(doc, "مفهوم ساده")
add_bullet_rtl(doc, "تفسیر آسان")

add_body_rtl(doc, "معایب:", bold=True)
add_bullet_rtl(doc, "مصرف حافظه زیاد")
add_bullet_rtl(doc, "پیاده‌سازی عملگرهای ژنتیکی دشوارتر")

add_heading_rtl(doc, "نمایش بردار عدد صحیح (Integer Vector)", level=2)
add_body_rtl(doc,
    "در این نمایش، هر خانه از بردار شماره استاد تخصیص‌یافته به همان دانشجو را نگهداری می‌کند.")

add_body_rtl(doc, "به عنوان مثال:")
add_code_block(doc, "[2, 1, 4, 3, 2, 5, ...]")

add_body_rtl(doc, "به این معنی که:")
add_bullet_rtl(doc, "Student1 → Professor2")
add_bullet_rtl(doc, "Student2 → Professor1")
add_bullet_rtl(doc, "Student3 → Professor4")
add_bullet_rtl(doc, "...")

add_heading_rtl(doc, "دلیل انتخاب این نمایش", level=2)
add_body_rtl(doc, "این نمایش برای مسئله حاضر مناسب‌تر است زیرا:")

reasons = [
    "حافظه بسیار کمی مصرف می‌کند.",
    "عملیات ارزیابی سریع‌تر انجام می‌شود.",
    "پیاده‌سازی Mutation و Crossover ساده‌تر است.",
    "در بسیاری از مسائل Assignment از همین نمایش استفاده شده است.",
    "با هر دو الگوریتم NSGA-II و Weighted Sum سازگار است.",
]
for item in reasons:
    add_bullet_rtl(doc, item)

add_body_rtl(doc,
    "بنابراین در این پروژه از Integer Vector Representation استفاده خواهد شد.", bold=True)

add_separator(doc)

# ============================================================
# SECTION 4: SOLUTION SPACE
# ============================================================
add_heading_rtl(doc, "۴. فضای جواب (Solution Space)", level=1)

add_body_rtl(doc,
    "هر جواب نشان‌دهنده یک تخصیص کامل از تمامی دانشجویان به اساتید است. "
    "اگر تعداد دانشجویان را n و تعداد اساتید را m در نظر بگیریم، اندازه فضای جواب تقریباً برابر خواهد بود با:")

add_formula(doc, "m^n")

add_body_rtl(doc,
    "به همین دلیل جستجوی کامل امکان‌پذیر نیست و استفاده از الگوریتم‌های فراابتکاری ضروری است.")

add_separator(doc)

# ============================================================
# SECTION 5: OBJECTIVE FUNCTIONS
# ============================================================
add_heading_rtl(doc, "۵. توابع هدف", level=1)

add_body_rtl(doc, "در این پروژه سه تابع هدف در نظر گرفته شده است.")

add_heading_rtl(doc, "هدف اول: بیشینه‌سازی رضایت دانشجویان", level=2)
add_body_rtl(doc,
    "هدف از این تابع، تخصیص هر دانشجو به یکی از اساتید مورد علاقه او است. "
    "هرچه استاد انتخاب‌شده در اولویت بالاتری قرار داشته باشد، مقدار این تابع افزایش می‌یابد.")

add_heading_rtl(doc, "هدف دوم: بیشینه‌سازی رضایت اساتید", level=2)
add_body_rtl(doc,
    "اساتید نیز نسبت به دانشجویان دارای ترجیحات مشخصی هستند که می‌تواند بر اساس کیفیت دانشجو، "
    "معدل، علایق پژوهشی و دانشگاه محل تحصیل قبلی تعیین شود. این تابع تلاش می‌کند اساتید تا حد "
    "امکان دانشجویان مطلوب خود را دریافت کنند.")

add_heading_rtl(doc, "هدف سوم: افزایش عدالت", level=2)
add_body_rtl(doc,
    "در بسیاری از تخصیص‌ها، ممکن است اکثر دانشجویان با کیفیت بالا به چند استاد خاص اختصاص داده شوند. "
    "هدف سوم تلاش می‌کند کیفیت دانشجویان بین اساتید به صورت متعادل توزیع شود. "
    "برای اندازه‌گیری عدالت می‌توان از معیارهایی مانند واریانس یا انحراف معیار میانگین Quality Score "
    "دانشجویان هر استاد استفاده کرد.")

add_separator(doc)

# ============================================================
# SECTION 6: CONSTRAINTS
# ============================================================
add_heading_rtl(doc, "۶. محدودیت‌ها", level=1)

add_body_rtl(doc, "در این مرحله دو محدودیت به صورت Soft Constraint در نظر گرفته شده‌اند.")

add_heading_rtl(doc, "ظرفیت اساتید", level=2)
add_body_rtl(doc,
    "هر استاد دارای ظرفیت مشخصی است. در حالت ایده‌آل این ظرفیت نباید نقض شود، اما در نسخه اولیه "
    "پروژه، برای افزایش انعطاف‌پذیری الگوریتم و کمک به فرآیند جستجو، نقض محدود ظرفیت مجاز بوده "
    "و به ازای آن جریمه در تابع برازندگی اعمال می‌شود.")

add_heading_rtl(doc, "تطابق گرایش", level=2)
add_body_rtl(doc,
    "ترجیح داده می‌شود دانشجو به استادی با گرایش مشابه اختصاص داده شود. با این حال، برای جلوگیری "
    "از محدود شدن بیش از حد فضای جستجو، در این مرحله عدم تطابق گرایش نیز به صورت محدودیت نرم "
    "مدل شده و تنها باعث اعمال جریمه خواهد شد.")

add_separator(doc)

# ============================================================
# SECTION 7: ALGORITHM OPERATORS
# ============================================================
add_heading_rtl(doc, "۷. عملگرهای الگوریتم", level=1)

add_heading_rtl(doc, "تولید جمعیت اولیه", level=2)
add_body_rtl(doc,
    "جمعیت اولیه به صورت تصادفی تولید می‌شود، اما تا حد امکان سعی می‌شود تخصیص‌های اولیه منطقی "
    "بوده و ظرفیت‌ها و گرایش‌ها تا حد امکان رعایت شوند.")

add_heading_rtl(doc, "عملگر Mutation", level=2)
add_body_rtl(doc,
    "برای این پروژه، Random Reset Mutation به عنوان عملگر اصلی انتخاب می‌شود. در این روش، استاد "
    "تخصیص‌یافته به یک دانشجو به صورت تصادفی تغییر داده می‌شود.")

add_body_rtl(doc, "دلایل انتخاب:", bold=True)
add_bullet_rtl(doc, "سادگی پیاده‌سازی")
add_bullet_rtl(doc, "ایجاد تنوع مناسب")
add_bullet_rtl(doc, "سازگاری کامل با نمایش برداری")
add_bullet_rtl(doc, "رایج بودن در مسائل تخصیص")

add_body_rtl(doc,
    "در کنار آن، Swap Mutation نیز می‌تواند به عنوان عملگر مکمل استفاده شود تا دو دانشجو استادهای "
    "خود را جابه‌جا کنند و تغییرات ملایم‌تری در جواب ایجاد شود.")

add_heading_rtl(doc, "عملگر Crossover", level=2)
add_body_rtl(doc,
    "پس از بررسی گزینه‌های مختلف، Uniform Crossover مناسب‌ترین انتخاب است.")

add_body_rtl(doc, "دلیل انتخاب:", bold=True)
add_bullet_rtl(doc, "تصمیم تخصیص هر دانشجو تا حد زیادی مستقل از سایر دانشجویان است.")
add_bullet_rtl(doc, "این عملگر تنوع بیشتری نسبت به One-Point و Two-Point ایجاد می‌کند.")
add_bullet_rtl(doc, "برای نمایش برداری مناسب است.")

add_body_rtl(doc,
    "عملگرهای PMX و Order Crossover برای این مسئله مناسب نیستند، زیرا برای مسائل ترتیبی مانند "
    "TSP طراحی شده‌اند و فرض می‌کنند هر مقدار تنها یک بار در کروموزوم ظاهر می‌شود؛ در حالی که "
    "در مسئله حاضر چندین دانشجو می‌توانند به یک استاد تخصیص یابند.")

add_heading_rtl(doc, "عملگر Selection", level=2)
add_body_rtl(doc, "در هر دو الگوریتم از Binary Tournament Selection استفاده می‌شود.")
add_bullet_rtl(doc, "در روش Weighted Sum انتخاب افراد بر اساس مقدار Fitness انجام می‌شود.")
add_bullet_rtl(doc, "در NSGA-II معیار انتخاب بر اساس Pareto Rank و Crowding Distance خواهد بود.")

add_heading_rtl(doc, "Repair Operator", level=2)
add_body_rtl(doc,
    "پس از اعمال Crossover یا Mutation ممکن است برخی تخصیص‌ها ظرفیت اساتید را بیش از حد نقض "
    "کنند یا کیفیت جواب کاهش یابد. به همین دلیل استفاده از یک Repair Operator پیشنهاد می‌شود تا "
    "جواب‌های نامعتبر یا بسیار نامطلوب اصلاح شوند و الگوریتم زمان خود را صرف بررسی جواب‌های "
    "غیرواقعی نکند.")

add_separator(doc)

# ============================================================
# SECTION 8: ALGORITHM DIFFERENCES
# ============================================================
add_heading_rtl(doc, "۸. تفاوت طراحی در Weighted Sum و NSGA-II", level=1)

add_body_rtl(doc,
    "هر دو الگوریتم از نمایش جواب، عملگرهای Mutation، Crossover، تولید جمعیت اولیه و Repair "
    "یکسان استفاده می‌کنند. تفاوت اصلی در نحوه ارزیابی جواب‌ها است.")

add_body_rtl(doc,
    "در روش Weighted Sum، تمامی توابع هدف با استفاده از ضرایب وزنی به یک مقدار Fitness تبدیل می‌شوند.")

add_body_rtl(doc,
    "در مقابل، در NSGA-II هر تابع هدف به صورت مستقل حفظ شده و انتخاب جواب‌ها بر اساس مفهوم "
    "Pareto Dominance و Crowding Distance انجام می‌شود.")

add_separator(doc)

# ============================================================
# SECTION 9: TERMINATION CRITERIA
# ============================================================
add_heading_rtl(doc, "۹. شرط پایان", level=1)

add_body_rtl(doc,
    "برای امکان مقایسه منصفانه دو الگوریتم، پیشنهاد می‌شود هر دو با شرایط یکسان اجرا شوند. "
    "شرط پایان می‌تواند رسیدن به تعداد مشخصی از نسل‌ها باشد (برای مثال ۲۰۰ نسل).")

add_body_rtl(doc,
    "در مطالعات بعدی، می‌توان معیارهایی مانند عدم بهبود در چند نسل متوالی یا محدودیت زمانی را نیز بررسی کرد.")

add_separator(doc)

# ============================================================
# CONCLUSION
# ============================================================
add_heading_rtl(doc, "جمع‌بندی", level=1)

add_body_rtl(doc,
    "در این مرحله، ساختار کلی مسئله، داده‌های مورد نیاز، نحوه نمایش جواب، فضای جستجو، توابع هدف، "
    "محدودیت‌ها و اجزای اصلی الگوریتم‌های فراابتکاری مشخص شده است. این تصمیم‌ها با توجه به ماهیت "
    "مسئله تخصیص اتخاذ شده‌اند تا ضمن حفظ سادگی پیاده‌سازی، امکان استفاده از هر دو روش Weighted "
    "Sum و NSGA-II بدون تغییر در ساختار اصلی مدل فراهم شود.")

add_body_rtl(doc,
    "در مراحل بعدی، بر اساس این طراحی، فایل‌های داده، مدل ریاضی و سپس پیاده‌سازی الگوریتم‌ها "
    "توسعه خواهند یافت.")

# ============================================================
# ADD PAGE NUMBERS
# ============================================================
def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

section = doc.sections[0]
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_number(footer)

# ============================================================
# SAVE
# ============================================================
output_path = r"D:\darsha\arshad\term2\algorithm\project\Project_Design_Document.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
