"""Generate comprehensive Persian project documentation."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Set RTL direction for Persian text."""
    pPr = paragraph._p.get_or_add_pPr()
    pBidi = OxmlElement("w:bidi")
    pPr.append(pBidi)

def set_font(run, font_name="B Nazanin", size=12, bold=False, color=None):
    """Set font for Persian text."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

def add_heading(doc, text, level=1):
    """Add RTL heading."""
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=18, bold=True, color=(31, 58, 95))
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_font(run, size=14, bold=True, color=(31, 58, 95))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        set_font(run, size=12, bold=True, color=(80, 80, 80))
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)

def add_body(doc, text, bold=False):
    """Add RTL body text."""
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_font(run, size=12, bold=bold)
    p.paragraph_format.line_spacing = 1.8
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    """Add RTL bullet point."""
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.line_spacing = 1.8
    p.paragraph_format.space_after = Pt(4)
    bullet_run = p.add_run("• ")
    set_font(bullet_run, size=12)
    text_run = p.add_run(text)
    set_font(text_run, size=12)

def create_document():
    """Create comprehensive Persian project document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "B Nazanin"
    font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "B Nazanin")

    # ========== COVER PAGE ==========
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("سند جامع پروژه")
    set_font(run, size=28, bold=True, color=(31, 58, 95))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("سیستم بهینه‌سازی تخصیص دانشجو به استاد راهنما")
    set_font(run, size=20, bold=True, color=(31, 58, 95))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("با استفاده از الگوریتم‌های ژنتیک چندهدفه")
    set_font(run, size=16, color=(100, 100, 100))

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("Weighted Sum GA و NSGA-II")
    set_font(run, size=14, color=(100, 100, 100))

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, "فهرست مطالب", level=1)

    toc_items = [
        "۱. معرفی پروژه",
        "۲. تعریف مسئله",
        "۳. توابع هدف",
        "۴. محدودیت‌ها",
        "۵. معماری نرم‌افزار",
        "۶. مدل‌های داده",
        "۷. الگوریتم Weighted Sum GA",
        "۸. الگوریتم NSGA-II",
        "۹. مقایسه الگوریتم‌ها",
        "۱۰. رابط کاربری وب",
        "۱۱. مجموعه داده‌ها",
        "۱۲. نتایج آزمایشات",
        "۱۳. جمع‌بندی",
    ]
    for item in toc_items:
        add_body(doc, item)

    doc.add_page_break()

    # ========== SECTION 1: INTRODUCTION ==========
    add_heading(doc, "۱. معرفی پروژه", level=1)

    add_body(doc, "این پروژه یک سیستم بهینه‌سازی چندهدفه برای تخصیص دانشجویان کارشناسی ارشد به اساتید راهنما است. هدف اصلی پیدا کردن بهترین تخصیص ممکن است که رضایت دانشجویان و اساتید را به حداکثر برساند و در عین حال توزیع عادلانه‌ای ایجاد کند.")

    add_body(doc, "این سیستم از دو الگوریتم فراابتکاری استفاده می‌کند:", bold=True)
    add_bullet(doc, "Weighted Sum Genetic Algorithm (GA): ترکیب اهداف چندگانه در یک تابع واحد با استفاده از وزن‌ها")
    add_bullet(doc, "NSGA-II: الگوریتم ژنتیک مرتب‌سازی نادومینان برای بهینه‌سازی چندهدفه واقعی")

    add_body(doc, "مزیت اصلی این رویکرد، جداسازی کامل لایه زیرساخت مشترک از لایه الگوریتم‌ها است که امکان توسعه و مقایسه آسان الگوریتم‌های مختلف را فراهم می‌کند.")

    # ========== SECTION 2: PROBLEM DEFINITION ==========
    add_heading(doc, "۲. تعریف مسئله", level=1)

    add_body(doc, "مسئله تخصیص دانشجو به استاد راهنما یک مسئله بهینه‌سازی ترکیبی (Combinatorial Optimization) است که در آن:")

    add_bullet(doc, "تعداد n دانشجو و m استاد وجود دارد")
    add_bullet(doc, "هر دانشجو باید دقیقاً به یک استاد تخصیص داده شود")
    add_bullet(doc, "هر استاد می‌تواند چندین دانشجو داشته باشد (با محدودیت ظرفیت)")
    add_bullet(doc, "دانشجویان و اساتید هر کدام لیست اولویت‌های خود را دارند")

    add_body(doc, "تعداد حالت‌های ممکن تخصیص برابر است با m^n که برای ۶۰ دانشجو و ۱۲ استاد برابر ۱۲^۶۰ است. این عدد بسیار بزرگ است و جستجوی کامل را غیرممکن می‌کند، بنابراین از الگوریتم‌های فراابتکاری استفاده می‌کنیم.")

    add_heading(doc, "نمایش جواب (Solution Representation)", level=2)

    add_body(doc, "در این پروژه از نمایش بردار عدد صحیح (Integer Vector) استفاده می‌شود. هر کروموزوم شامل یک بردار است که هر خانه آن شماره استاد تخصیص‌یافته به یک دانشجو را نگه می‌دارد.")

    add_body(doc, "مثال: [2, 1, 4, 3, 2, 5, ...] به این معنی که:")
    add_bullet(doc, "دانشجوی ۱ به استاد ۲ تخصیص داده شده")
    add_bullet(doc, "دانشجوی ۲ به استاد ۱ تخصیص داده شده")
    add_bullet(doc, "دانشجوی ۳ به استاد ۴ تخصیص داده شده")
    add_bullet(doc, "...")

    add_body(doc, "دلیل انتخاب این نمایش: سادگی پیاده‌سازی، مصرف حافظه کم، و سازگاری با هر دو الگوریتم.")

    # ========== SECTION 3: OBJECTIVE FUNCTIONS ==========
    add_heading(doc, "۳. توابع هدف", level=1)

    add_body(doc, "این پروژه سه تابع هدف دارد که باید بهینه‌سازی شوند:")

    add_heading(doc, "هدف اول: رضایت دانشجویان", level=2)
    add_body(doc, "این تابع تلاش می‌کند هر دانشجو را به یکی از اساتید مورد علاقه‌اش تخصیص دهد. هرچه استاد انتخاب‌شده در اولویت بالاتری باشد، رضایت بیشتر است.")
    add_body(doc, "فرمول: مجموع (1/رتبه) برای تمام دانشجویان، نرمال‌شده به بازه [0, 1]")

    add_heading(doc, "هدف دوم: رضایت اساتید", level=2)
    add_body(doc, "اساتید نیز نسبت به دانشجویان اولویت‌هایی دارند. این تابع تلاش می‌کند اساتید دانشجویان مطلوب خود را دریافت کنند.")
    add_body(doc, "فرمول: مجموع (1/رتبه) برای تمام تخصیص‌ها، نرمال‌شده به بازه [0, 1]")

    add_heading(doc, "هدف سوم: عدالت", level=2)
    add_body(doc, "توزیع دانشجویان با کیفیت بین اساتید باید عادلانه باشد. از واریانس میانگین کیفیت دانشجویان هر استاد استفاده می‌شود.")
    add_body(doc, "فرمول: واریانس میانگین Quality Score دانشجویان هر استاد، نرمال‌شده به بازه [0, 1]")

    # ========== SECTION 4: CONSTRAINTS ==========
    add_heading(doc, "۴. محدودیت‌ها", level=1)

    add_body(doc, "دو محدودیت اصلی در این مسئله وجود دارد:")

    add_heading(doc, "محدودیت اول: ظرفیت اساتید", level=2)
    add_body(doc, "هر استاد حداکثر تعداد مشخصی دانشجو می‌تواند داشته باشد. در نسخه اولیه این محدودیت نرم (Soft Constraint) است و نقض آن با جریمه مجازات می‌شود.")
    add_body(doc, "امکان تنظیم جریمه به صورت 'inf' برای تبدیل به محدودیت سخت وجود دارد.")

    add_heading(doc, "محدودیت دوم: تطابق گرایش", level=2)
    add_body(doc, "ترجیح داده می‌شود دانشجو به استادی با گرایش مشابه تخصیص داده شود. عدم تطابق نیز با جریمه مجازات می‌شود.")
    add_body(doc, "گرایش‌های موجود: هوش مصنوعی، مهندسی نرم‌افزار، شبکه‌های کامپیوتری، معماری کامپیوتر")

    # ========== SECTION 5: SOFTWARE ARCHITECTURE ==========
    add_heading(doc, "۵. معماری نرم‌افزار", level=1)

    add_body(doc, "معماری نرم‌افزار بر اساس اصول SOLID طراحی شده و شامل سه لایه اصلی است:")

    add_heading(doc, "لایه زیرساخت مشترک", level=2)
    add_body(doc, "این لایه شامل کلاس‌هایی است که توسط هر دو الگوریتم استفاده می‌شوند:")
    add_bullet(doc, "مدل‌های داده (Student, Professor, Chromosome, Population)")
    add_bullet(doc, "ارزیاب‌ها (ObjectiveEvaluator, ConstraintEvaluator)")
    add_bullet(doc, "عملگرها (Mutation, Crossover)")
    add_bullet(doc, "مقداردهی اولیه جمعیت")
    add_bullet(doc, "تعمیر جواب‌ها (Repair Operator)")

    add_heading(doc, "لایه الگوریتم‌ها", level=2)
    add_body(doc, "هر الگوریتم ماژول جداگانه‌ای دارد:")
    add_bullet(doc, "Weighted Sum GA: محاسبه فیتنس وزنی، انتخاب تورنمنت، جایگزینی نخبه")
    add_bullet(doc, "NSGA-II: مرتب‌سازی نادومینان، فاصلهcrowding، انتخاب بر اساس رتبه")

    add_heading(doc, "لایه رابط کاربری", level=2)
    add_body(doc, "رابط وب با FastAPI و HTML ساده برای اجرای الگوریتم‌ها و مشاهده نتایج.")

    # ========== SECTION 6: DATA MODELS ==========
    add_heading(doc, "۶. مدل‌های داده", level=1)

    add_heading(doc, "مدل دانشجو (Student)", level=2)
    add_bullet(doc, "student_id: شناسه یکتا")
    add_bullet(doc, "name: نام دانشجو")
    add_bullet(doc, "field: گرایش تخصصی")
    add_bullet(doc, "quality_score: امتیاز کیفیت (محاسبه‌شده از معدل و کیفیت دانشگاه)")
    add_bullet(doc, "preference_list: لیست اولویت اساتید")

    add_heading(doc, "مدل استاد (Professor)", level=2)
    add_bullet(doc, "professor_id: شناسه یکتا")
    add_bullet(doc, "name: نام استاد")
    add_bullet(doc, "field: گرایش تخصصی")
    add_bullet(doc, "capacity: حداکثر ظرفیت")
    add_bullet(doc, "preference_list: لیست اولویت دانشجویان")

    add_heading(doc, "مدل کروموزوم (Chromosome)", level=2)
    add_body(doc, "کروموزوم نماینده یک جواب کامل (تمام تخصیص‌ها) است:")
    add_bullet(doc, "genes: برداری از شناسه اساتید تخصیص‌یافته")
    add_bullet(doc, "objectives: مقادیر سه هدف")
    add_bullet(doc, "constraints: مقدار نقض محدودیت‌ها")
    add_bullet(doc, "fitness: مقدار فیتنس (در Weighted Sum)")
    add_bullet(doc, "metadata: اطلاعات اضافی (رتبه و فاصلهcrowding در NSGA-II)")

    add_heading(doc, "مدل جمعیت (Population)", level=2)
    add_body(doc, "جمعیت مجموعه‌ای از کروموزوم‌ها است که در هر نسل تکامل می‌یابد.")

    # ========== SECTION 7: WEIGHTED SUM GA ==========
    add_heading(doc, "۷. الگوریتم Weighted Sum GA", level=1)

    add_body(doc, "این الگوریتم اهداف چندگانه را با استفاده از وزن‌ها در یک مقدار فیتنس واحد ترکیب می‌کند.")

    add_heading(doc, "فرمول فیتنس", level=2)
    add_body(doc, "fitness = w1 × رضایت_دانشجو + w2 × رضایت_استاد + w3 × (1 - عدالت) - جریمه_محدودیت")

    add_body(doc, "پارامترهای قابل تنظیم:")
    add_bullet(doc, "وزن‌های هدف: w1, w2, w3 (مجموع باید ۱ باشد)")
    add_bullet(doc, "جریمه محدودیت ظرفیت (قابل تنظیم به inf)")
    add_bullet(doc, "جریمه عدم تطابق گرایش (قابل تنظیم به inf)")
    add_bullet(doc, "اندازه تورنمنت: تعداد افراد در هر تورنمنت انتخاب")
    add_bullet(doc, "اندازه نخبگان: تعداد بهترین افراد حفظ‌شده در هر نسل")

    add_heading(doc, "مراحل الگوریتم", level=2)
    add_bullet(doc, "۱. مقداردهی اولیه جمعیت")
    add_bullet(doc, "۲. ارزیابی تمام کروموزوم‌ها")
    add_bullet(doc, "۳. تکرار تا حداکثر نسل‌ها:")
    add_bullet(doc, "   الف. انتخاب والدین (تورنمنت)")
    add_bullet(doc, "   ب. اعمال crossover")
    add_bullet(doc, "   ج. اعمال mutation")
    add_bullet(doc, "   د. تعمیر جواب‌ها")
    add_bullet(doc, "   ه. ارزیابی فرزندان")
    add_bullet(doc, "   و. جایگزینی با حفظ نخبگان")
    add_bullet(doc, "۴. بازگرداندن بهترین جواب")

    # ========== SECTION 8: NSGA-II ==========
    add_heading(doc, "۸. الگوریتم NSGA-II", level=1)

    add_body(doc, "NSGA-II یک الگوریتم ژنتیک چندهدفه واقعی است که به جای ترکیب اهداف، مجموعه‌ای از جواب‌های بهینه پارتو را حفظ می‌کند.")

    add_heading(doc, "مرتب‌سازی نادومینان (Fast Non-Dominated Sorting)", level=2)
    add_body(doc, "این الگوریتم کروموزوم‌ها را بر اساس مفهوم تسلط پارتو رتبه‌بندی می‌کند:")
    add_bullet(doc, "رتبه ۰: جواب‌هایی که هیچ جواب دیگری بر آنها تسلط ندارد (بهترین)")
    add_bullet(doc, "رتبه ۱: جواب‌هایی که فقط توسط رتبه ۰ تسلط دارند")
    add_bullet(doc, "... و الی آخر")

    add_heading(doc, "فاصله Crowding", level=2)
    add_body(doc, "برای حفظ تنوع در جمعیت، فاصله crowding محاسبه می‌شود. جواب‌هایی که در مناطق کمتر متراکم هستند، فاصله بیشتری دارند و احتمال انتخاب آنها بالاتر است.")

    add_heading(doc, "انتخاب تورنمنت NSGA-II", level=2)
    add_body(doc, "در NSGA-II، انتخاب تورنمنت بر اساس:")
    add_bullet(doc, "اول: رتبه پارتو (رتبه کمتر بهتر)")
    add_bullet(doc, "دوم: فاصله crowding (فاصله بیشتر بهتر)")

    add_heading(doc, "مراحل الگوریتم", level=2)
    add_bullet(doc, "۱. مقداردهی اولیه جمعیت P")
    add_bullet(doc, "۲. ارزیابی تمام کروموزوم‌ها")
    add_bullet(doc, "۳. تکرار تا حداکثر نسل‌ها:")
    add_bullet(doc, "   الف. تولید فرزندان Q با انتخاب تورنمنت")
    add_bullet(doc, "   ب. اعمال crossover و mutation")
    add_bullet(doc, "   ج. تعمیر و ارزیابی فرزندان")
    add_bullet(doc, "   د. ادغام R = P + Q")
    add_bullet(doc, "   ه. مرتب‌سازی نادومینان")
    add_bullet(doc, "   و. محاسبه فاصله crowding")
    add_bullet(doc, "   ز. پر کردن نسل بعدی از بهترین فرانت‌ها")
    add_bullet(doc, "۴. استخراج فرانت پارتو نهایی")

    # ========== SECTION 9: ALGORITHM COMPARISON ==========
    add_heading(doc, "۹. مقایسه الگوریتم‌ها", level=1)

    add_heading(doc, "تفاوت‌های اصلی", level=2)

    add_body(doc, "Weighted Sum GA:")
    add_bullet(doc, "اهداف را در یک مقدار فیتنس واحد ترکیب می‌کند")
    add_bullet(doc, "یک جواب بهینه واحد برمی‌گرداند")
    add_bullet(doc, "نیاز به انتخاب وزن‌ها دارد")
    add_bullet(doc, "سریع‌تر اجرا می‌شود")

    add_body(doc, "NSGA-II:")
    add_bullet(doc, "اهداف را جداگانه حفظ می‌کند")
    add_bullet(doc, "مجموعه‌ای از جواب‌های بهینه پارتو برمی‌گرداند")
    add_bullet(doc, "نیازی به وزن‌ها ندارد")
    add_bullet(doc, "تنوع بیشتری در جواب‌ها ایجاد می‌کند")

    add_heading(doc, "معیارهای مقایسه", level=2)
    add_bullet(doc, "رضایت دانشجویان (میانه، چارک اول، چارک سوم)")
    add_bullet(doc, "رضایت اساتید")
    add_bullet(doc, "عدالت")
    add_bullet(doc, "زمان اجرا")
    add_bullet(doc, "تعداد جواب‌های بهینه پارتو (فقط NSGA-II)")

    add_heading(doc, "آمار استنباطی", level=2)
    add_body(doc, "برای مقایسه آماری از آزمون Mann-Whitney U Test استفاده می‌شود:")
    add_bullet(doc, "سطح معناداری: α = 0.05")
    add_bullet(doc, "تعداد اجراها: ۳۰ بار برای هر الگوریتم")
    add_bullet(doc, "گزارش: مقدار U، مقدار p، معنادار بودن تفاوت")

    # ========== SECTION 10: WEB INTERFACE ==========
    add_heading(doc, "۱۰. رابط کاربری وب", level=1)

    add_body(doc, "رابط کاربری وب با FastAPI و HTML ساده پیاده‌سازی شده و شامل بخش‌های زیر است:")

    add_heading(doc, "بخش مدیریت داده‌ها", level=2)
    add_bullet(doc, "انتخاب مجموعه داده از پیش تعریف‌شده (کوچک، متوسط، بزرگ)")
    add_bullet(doc, "آپلود مجموعه داده سفارشی")
    add_bullet(doc, "پیش‌نمایش محتوای داده‌ها")

    add_heading(doc, "بخش اجرای تکی", level=2)
    add_bullet(doc, "انتخاب الگوریتم (Weighted Sum یا NSGA-II)")
    add_bullet(doc, "تنظیم پارامترهای الگوریتم")
    add_bullet(doc, "تنظیم وزن‌ها و جریمه‌ها (فقط Weighted Sum)")
    add_bullet(doc, "مشاهده نتایج و جدول تخصیص‌ها")

    add_heading(doc, "بخش مقایسه", level=2)
    add_bullet(doc, "اجرا هر دو الگوریتم چندین بار")
    add_bullet(doc, "گزارش آماری (میانه، چارک‌ها)")
    add_bullet(doc, "آزمون Mann-Whitney U")
    add_bullet(doc, "جدول مقایسه‌ای")

    add_heading(doc, "پارامترهای قابل تنظیم", level=2)
    add_bullet(doc, "اندازه جمعیت")
    add_bullet(doc, "حداکثر نسل‌ها")
    add_bullet(doc, "احتمال crossover")
    add_bullet(doc, "احتمال mutation")
    add_bullet(doc, "اندازه تورنمنت")
    add_bullet(doc, "اندازه نخبگان")
    add_bullet(doc, "وزن‌های هدف (Weighted Sum)")
    add_bullet(doc, "جریمه‌های محدودیت (قابل تنظیم به inf)")

    # ========== SECTION 11: DATASETS ==========
    add_heading(doc, "۱۱. مجموعه داده‌ها", level=1)

    add_body(doc, "سیستم دارای مجموعه داده‌های از پیش تعریف‌شده است:")

    add_heading(doc, "مجموعه داده کوچک", level=2)
    add_bullet(doc, "۲۰ دانشجو")
    add_bullet(doc, "۸ استاد")
    add_bullet(doc, "مناسب برای آزمایش سریع")

    add_heading(doc, "مجموعه داده متوسط", level=2)
    add_bullet(doc, "۶۰ دانشجو")
    add_bullet(doc, "۱۲ استاد")
    add_bullet(doc, "مناسب برای مقایسه الگوریتم‌ها")

    add_heading(doc, "مجموعه داده بزرگ", level=2)
    add_bullet(doc, "۱۵۰ دانشجو")
    add_bullet(doc, "۲۸ استاد")
    add_bullet(doc, "مناسب برای آزمایش مقیاس‌پذیری")

    add_heading(doc, "ساختار فایل‌ها", level=2)
    add_bullet(doc, "Students.xlsx: اطلاعات دانشجویان")
    add_bullet(doc, "Professors.xlsx: اطلاعات اساتید")
    add_bullet(doc, "StudentPreferences.xlsx: اولویت‌های دانشجویان")
    add_bullet(doc, "ProfessorPreferences.xlsx: اولویت‌های اساتید")
    add_bullet(doc, "Universities.xlsx: اطلاعات دانشگاه‌ها")
    add_bullet(doc, "Fields.xlsx: گرایش‌های تخصصی")

    # ========== SECTION 12: RESULTS ==========
    add_heading(doc, "۱۲. نتایج آزمایشات", level=1)

    add_body(doc, "نتایج اجرای آزمایشی الگوریتم‌ها:")

    add_heading(doc, "Weighted Sum GA", level=2)
    add_bullet(doc, "رضایت دانشجویان: ۸۷.۵٪")
    add_bullet(doc, "رضایت اساتید: ۱۹.۰٪")
    add_bullet(doc, "عدالت: ۰.۰۱۷۱")
    add_bullet(doc, "زمان اجرا: ۱.۲۴ ثانیه")

    add_heading(doc, "NSGA-II", level=2)
    add_bullet(doc, "رضایت دانشجویان: ۶۶.۶٪")
    add_bullet(doc, "رضایت اساتید: ۲۰.۴٪")
    add_bullet(doc, "عدالت: ۰.۰۰۵۸")
    add_bullet(doc, "تعداد جواب‌های پارتو: ۱۵")
    add_bullet(doc, "زمان اجرا: ۳.۵۲ ثانیه")

    add_body(doc, "نتیجه‌گیری: Weighted Sum GA در رضایت دانشجویان بهتر عمل می‌کند، در حالی که NSGA-II عدالت بیشتری ایجاد می‌کند و مجموعه‌ای متنوع از جواب‌ها ارائه می‌دهد.")

    # ========== SECTION 13: CONCLUSION ==========
    add_heading(doc, "۱۳. جمع‌بندی", level=1)

    add_body(doc, "در این پروژه یک سیستم جامع برای بهینه‌سازی تخصیص دانشجو به استاد راهنما پیاده‌سازی شد. ویژگی‌های اصلی سیستم:")

    add_bullet(doc, "معماری ماژولار و قابل توسعه")
    add_bullet(doc, "پیاده‌سازی دو الگوریتم چندهدفه (Weighted Sum و NSGA-II)")
    add_bullet(doc, "امکان مقایسه آماری الگوریتم‌ها")
    add_bullet(doc, "رابط کاربری وب ساده و کاربردی")
    add_bullet(doc, "پشتیبانی از مجموعه داده‌های مختلف")
    add_bullet(doc, "امکان تنظیم تمام پارامترها")

    add_body(doc, "این سیستم قابلیت استفاده در پایان‌نامه کارشناسی ارشد را دارد و می‌تواند برای مقایسه الگوریتم‌های بهینه‌سازی چندهدفه استفاده شود.")

    # Save document
    output_path = r"D:\darsha\arshad\term2\algorithm\project\Project_Document_FA.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_document()
